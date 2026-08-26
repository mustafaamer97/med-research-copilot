from docx import Document


def export_to_docx(
    content,
    title,
    output_file
):

    doc = Document()

    doc.add_heading(
        title,
        level=1
    )

    for line in content.split("\n"):

        doc.add_paragraph(
            line
        )

    doc.save(
        output_file
    )

    return output_file
